from pathlib import Path
from typing import Protocol

import docx
import pypdf

from buscador_de_vaga.resume.exceptions import (
    EmptyDocumentError,
    ResumeReadError,
    UnreadablePdfError,
    UnsupportedFileFormatError,
)
from buscador_de_vaga.resume.models import RawResumeText


class ResumeReader(Protocol):
    """Interface (Protocol) para leitores técnicos de arquivos de currículo."""

    def read(self, file_path: Path | str) -> RawResumeText:
        ...


class PdfResumeReader:
    """Leitor de arquivos PDF utilizando pypdf."""

    def read(self, file_path: Path | str) -> RawResumeText:
        path = Path(file_path)
        if not path.is_file():
            raise ResumeReadError(f"Arquivo não encontrado: {path}")
        if path.stat().st_size == 0:
            raise EmptyDocumentError(f"O arquivo {path.name} está vazio (0 bytes).")

        try:
            reader = pypdf.PdfReader(str(path))
            page_count = len(reader.pages)
            if page_count == 0:
                raise EmptyDocumentError(f"O PDF {path.name} não possui páginas.")

            extracted_pages: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                extracted_pages.append(page_text)

            full_text = "\n".join(extracted_pages)
            if not full_text.strip():
                raise UnreadablePdfError(
                    f"O PDF {path.name} possui {page_count} página(s), mas não possui "
                    "camada de texto detectável. OCR ainda não é suportado."
                )

            return RawResumeText(
                source_file=str(path),
                text=full_text,
                page_count=page_count,
            )
        except (UnreadablePdfError, EmptyDocumentError, ResumeReadError):
            raise
        except Exception as exc:
            raise ResumeReadError(f"Erro ao ler arquivo PDF {path.name}: {exc}") from exc


class DocxResumeReader:
    """Leitor de arquivos DOCX utilizando python-docx."""

    def read(self, file_path: Path | str) -> RawResumeText:
        path = Path(file_path)
        if not path.is_file():
            raise ResumeReadError(f"Arquivo não encontrado: {path}")
        if path.stat().st_size == 0:
            raise EmptyDocumentError(f"O arquivo {path.name} está vazio (0 bytes).")

        try:
            doc = docx.Document(str(path))
            lines: list[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    lines.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))

            full_text = "\n".join(lines)
            if not full_text.strip():
                raise EmptyDocumentError(f"O arquivo DOCX {path.name} não contém texto extraível.")

            return RawResumeText(
                source_file=str(path),
                text=full_text,
                page_count=1,
            )
        except (EmptyDocumentError, ResumeReadError):
            raise
        except Exception as exc:
            raise ResumeReadError(f"Erro ao ler arquivo DOCX {path.name}: {exc}") from exc


def read_resume(file_path: Path | str) -> RawResumeText:
    """Função utilitária que seleciona o leitor apropriado com base na extensão do arquivo."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return PdfResumeReader().read(path)
    elif ext == ".docx":
        return DocxResumeReader().read(path)
    else:
        raise UnsupportedFileFormatError(
            f"Formato de arquivo '{ext}' não suportado. Por favor, utilize arquivos .pdf ou .docx."
        )
