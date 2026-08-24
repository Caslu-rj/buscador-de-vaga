from pathlib import Path

import docx
import pytest
from pypdf import PdfWriter

from buscador_de_vaga.resume.exceptions import (
    EmptyDocumentError,
    ResumeReadError,
    UnreadablePdfError,
    UnsupportedFileFormatError,
)
from buscador_de_vaga.resume.models import RawResumeText
from buscador_de_vaga.resume.reader import (
    DocxResumeReader,
    PdfResumeReader,
    read_resume,
)


def create_synthetic_pdf(path: Path, text: str) -> Path:
    """Gera um PDF sintético com a string de texto fornecida."""
    pdf_content = (
        "%PDF-1.4\n"
        "1 0 obj\n"
        "<< /Type /Catalog /Pages 2 0 R >>\n"
        "endobj\n"
        "2 0 obj\n"
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>\n"
        "endobj\n"
        "3 0 obj\n"
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        "/Resources << /Font << /F1 5 0 R >> >> >>\n"
        "endobj\n"
        "4 0 obj\n"
        f"<< /Length {len(f'BT /F1 12 Tf 100 700 Td ({text}) Tj ET')} >>\n"
        "stream\n"
        f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET\n"
        "endstream\n"
        "endobj\n"
        "5 0 obj\n"
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        "endobj\n"
        "xref\n"
        "0 6\n"
        "0000000000 65535 f \n"
        "0000000009 00000 n \n"
        "0000000058 00000 n \n"
        "0000000115 00000 n \n"
        "0000000246 00000 n \n"
        "0000000360 00000 n \n"
        "trailer\n"
        "<< /Size 6 /Root 1 0 R >>\n"
        "startxref\n"
        "440\n"
        "%%EOF\n"
    )
    path.write_bytes(pdf_content.encode("latin1"))
    return path


def create_blank_pdf(path: Path) -> Path:
    """Gera um PDF válido com 1 página porém sem qualquer texto (escaneado/em branco)."""
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as f:
        writer.write(f)
    return path


def create_synthetic_docx(
    path: Path,
    paragraphs: list[str],
    table_data: list[list[str]] | None = None,
) -> Path:
    """Gera um arquivo DOCX sintético com parágrafos e tabela opcional."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for row_idx, row_values in enumerate(table_data):
            for col_idx, cell_value in enumerate(row_values):
                table.cell(row_idx, col_idx).paragraphs[0].text = cell_value
    doc.save(str(path))
    return path


# --- TESTES DE PDF ---


def test_pdf_resume_reader_success(tmp_path: Path) -> None:
    pdf_path = create_synthetic_pdf(tmp_path / "curriculo.pdf", "Ana Silva - Desenvolvedora Python")
    reader = PdfResumeReader()
    result = reader.read(pdf_path)

    assert isinstance(result, RawResumeText)
    assert result.source_file == str(pdf_path)
    assert "Ana Silva" in result.text
    assert result.page_count == 1


def test_pdf_resume_reader_short_text_success(tmp_path: Path) -> None:
    """PDFs curtos com texto válido devem ser aceitos sem erro arbitrário de tamanho."""
    pdf_path = create_synthetic_pdf(tmp_path / "curto.pdf", "Ana")
    reader = PdfResumeReader()
    result = reader.read(pdf_path)

    assert result.text.strip() == "Ana"
    assert result.page_count == 1


def test_pdf_resume_reader_unreadable_scanned_pdf(tmp_path: Path) -> None:
    """PDFs com páginas mas sem camada de texto (escaneados) devem lançar UnreadablePdfError."""
    pdf_path = create_blank_pdf(tmp_path / "escaneado.pdf")
    reader = PdfResumeReader()

    with pytest.raises(UnreadablePdfError, match=r"camada de texto detectável"):
        reader.read(pdf_path)


def test_pdf_resume_reader_empty_0_byte_file(tmp_path: Path) -> None:
    """Arquivos com 0 bytes devem lançar EmptyDocumentError."""
    pdf_path = tmp_path / "vazio.pdf"
    pdf_path.write_bytes(b"")
    reader = PdfResumeReader()

    with pytest.raises(EmptyDocumentError):
        reader.read(pdf_path)


def test_pdf_resume_reader_corrupted_file(tmp_path: Path) -> None:
    """Arquivos corrompidos devem lançar ResumeReadError."""
    pdf_path = tmp_path / "corrompido.pdf"
    pdf_path.write_bytes(b"NAO_E_UM_PDF_VALIDO_12345")
    reader = PdfResumeReader()

    with pytest.raises(ResumeReadError):
        reader.read(pdf_path)


# --- TESTES DE DOCX ---


def test_docx_resume_reader_success_paragraphs(tmp_path: Path) -> None:
    docx_path = create_synthetic_docx(
        tmp_path / "curriculo.docx",
        paragraphs=["Carlos Eduardo", "Engenheiro de Dados", "Habilidades: Python, SQL"],
    )
    reader = DocxResumeReader()
    result = reader.read(docx_path)

    assert isinstance(result, RawResumeText)
    assert result.source_file == str(docx_path)
    assert "Carlos Eduardo" in result.text
    assert "Habilidades: Python, SQL" in result.text


def test_docx_resume_reader_success_with_tables(tmp_path: Path) -> None:
    docx_path = create_synthetic_docx(
        tmp_path / "curriculo_tabela.docx",
        paragraphs=["Experiência Profissional"],
        table_data=[["Empresa Tech", "Desenvolvedor Backend"], ["2020 - 2024", "Python e FastAPI"]],
    )
    reader = DocxResumeReader()
    result = reader.read(docx_path)

    assert "Experiência Profissional" in result.text
    assert "Empresa Tech" in result.text
    assert "Desenvolvedor Backend" in result.text


def test_docx_resume_reader_empty_document(tmp_path: Path) -> None:
    """DOCX sem qualquer texto em parágrafos ou tabelas deve lançar EmptyDocumentError."""
    docx_path = create_synthetic_docx(tmp_path / "sem_texto.docx", paragraphs=[])
    reader = DocxResumeReader()

    with pytest.raises(EmptyDocumentError):
        reader.read(docx_path)


# --- TESTES DISPATCHER / FORMATO NAO SUPORTADO ---


def test_read_resume_dispatcher_pdf(tmp_path: Path) -> None:
    pdf_path = create_synthetic_pdf(tmp_path / "curriculo.pdf", "Maria Souza")
    result = read_resume(pdf_path)
    assert "Maria Souza" in result.text


def test_read_resume_dispatcher_docx(tmp_path: Path) -> None:
    docx_path = create_synthetic_docx(tmp_path / "curriculo.docx", paragraphs=["Maria Souza"])
    result = read_resume(docx_path)
    assert "Maria Souza" in result.text


def test_read_resume_unsupported_extension(tmp_path: Path) -> None:
    txt_path = tmp_path / "curriculo.txt"
    txt_path.write_text("Maria Souza", encoding="utf-8")

    with pytest.raises(UnsupportedFileFormatError, match=r"não suportado"):
        read_resume(txt_path)


def test_read_resume_file_not_found(tmp_path: Path) -> None:
    non_existent = tmp_path / "nao_existe.pdf"
    with pytest.raises(ResumeReadError):
        read_resume(non_existent)
